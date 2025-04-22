from docx import Document

def export_to_docx(filename, narrative, model_outputs):
    doc = Document()
    doc.add_heading('Political Risk Assessment Report', 0)
    doc.add_paragraph(narrative)

    for model in model_outputs:
        doc.add_heading(model['model'] + ' Model', level=2)
        doc.add_paragraph(f"Score: {model['score']}")
        doc.add_paragraph("Macro Risk: " + ', '.join(model['macro']))
        doc.add_paragraph("Micro Risk: " + ', '.join(model['micro']))
        doc.add_paragraph("Sovereign Risk: " + ', '.join(model['sovereign']))
        doc.add_paragraph("Recommendations:")
        doc.add_paragraph(model['recommendations'])

    doc.save(filename)
    return filename